"""
Reproducibly generates template.docx — minimal Jinja-aware Word template
for the docxtpl spike (pipeline-v6 Week 1).

Run:  uv run python make-template.py
"""
from docx import Document
from docx.shared import Pt


def main() -> None:
    doc = Document()

    # Title + author
    title = doc.add_paragraph()
    run = title.add_run("{{ title }}")
    run.bold = True
    run.font.size = Pt(20)

    author = doc.add_paragraph("Автор: {{ author }}")
    author.runs[0].font.size = Pt(12)

    # Loop over chapters via docxtpl block tags inside paragraphs
    doc.add_paragraph("{%p for ch in chapters %}")
    doc.add_heading("Глава {{ loop.index }}. {{ ch.title }}", level=1)
    doc.add_paragraph("{{ ch.body }}")
    doc.add_paragraph("{%p endfor %}")

    # Table with row-loop. docxtpl {%tr %} needs the open and close
    # tags in SEPARATE rows, otherwise both tags get consumed when
    # docxtpl deletes the containing <w:tr>.
    doc.add_heading("Таблица результатов", level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Параметр"
    hdr[2].text = "Значение"

    # Row 1: open tag
    table.rows[1].cells[0].text = "{%tr for row in table.rows %}"
    # Row 2: repeating content
    body = table.rows[2].cells
    body[0].text = "{{ loop.index }}"
    body[1].text = "{{ row.name }}"
    body[2].text = "{{ row.value }}"
    # Row 3: close tag
    table.rows[3].cells[0].text = "{%tr endfor %}"

    # Bibliography
    doc.add_heading("Список литературы", level=1)
    doc.add_paragraph("{%p for src in bibliography %}")
    doc.add_paragraph("{{ loop.index }}. {{ src }}")
    doc.add_paragraph("{%p endfor %}")

    doc.save("template.docx")
    print("template.docx written")


if __name__ == "__main__":
    main()
