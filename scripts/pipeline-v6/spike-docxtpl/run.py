"""
Pipeline-v6 spike: render template.docx via docxtpl with sample context,
inspect result with python-docx, print metrics.
"""
import os
import time

from docx import Document
from docxtpl import DocxTemplate


CONTEXT = {
    "title": "Цифровая трансформация ретейла",
    "author": "Иванов И.И.",
    "chapters": [
        {
            "title": "Введение",
            "body": (
                "Современный ретейл переживает фундаментальные изменения "
                "под влиянием цифровых технологий. В данной работе "
                "рассматриваются ключевые направления трансформации."
            ),
        },
        {
            "title": "Методология",
            "body": (
                "Исследование основано на сравнительном анализе "
                "open-source и проприетарных платформ электронной коммерции."
            ),
        },
    ],
    "table": {
        "rows": [
            {"name": "Конверсия", "value": "3.2%"},
            {"name": "AOV", "value": "1450 ₽"},
            {"name": "Retention D30", "value": "27%"},
        ],
    },
    "bibliography": [
        "Котлер Ф. Маркетинг 5.0. — М.: Эксмо, 2021.",
        "Иванов А.Б. Цифровой ретейл. — СПб.: Питер, 2023.",
    ],
}


def main() -> None:
    t0 = time.perf_counter()
    tpl = DocxTemplate("template.docx")
    tpl.render(CONTEXT)
    tpl.save("out.docx")
    elapsed_ms = (time.perf_counter() - t0) * 1000

    size = os.path.getsize("out.docx")
    doc = Document("out.docx")
    headings = sum(1 for p in doc.paragraphs if (p.style.name or "").startswith("Heading"))
    tables = len(doc.tables)
    paragraphs = len(doc.paragraphs)
    table_rows = sum(len(t.rows) for t in doc.tables)

    print("== docxtpl spike ==")
    print(f"render+save: {elapsed_ms:.1f} ms")
    print(f"out.docx size: {size} bytes")
    print(f"paragraphs: {paragraphs}")
    print(f"headings: {headings}")
    print(f"tables: {tables} (total rows: {table_rows})")


if __name__ == "__main__":
    main()
